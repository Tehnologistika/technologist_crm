from fastapi import FastAPI, HTTPException, UploadFile, File
from fastapi.middleware.cors import CORSMiddleware
from fastapi.staticfiles import StaticFiles
from fastapi.responses import FileResponse
from pydantic import BaseModel
from typing import Optional
import re
import requests
from datetime import datetime
from datetime import datetime, timedelta
from sqlalchemy import func, select
import json, os
try:
    from dotenv import load_dotenv
    load_dotenv()
except ModuleNotFoundError:
    print("python-dotenv not installed — skipping .env loading")
import secrets
from pathlib import Path 
BASE_DIR = Path(__file__).resolve().parent
INVITES_FILE = BASE_DIR / "invites.json"
import warnings
warnings.filterwarnings("ignore", category=DeprecationWarning)

from telegram import Update
from telegram.ext import ApplicationBuilder, CommandHandler, ContextTypes
from modules.wizards.close import register_close_conversation
import threading

# --- PDF act generation ---
from docx import Document
from docxtpl import DocxTemplate
import subprocess, tempfile, shutil
ACT_TEMPL_CUST = BASE_DIR / "template_cust.docx"
ACT_TEMPL_EXEC = BASE_DIR / "template_exec.docx"
ACTS_DIR = BASE_DIR / "acts"
ACTS_DIR.mkdir(exist_ok=True)
BASE_URL = "http://147.45.232.245:8000"   # ← change to your domain if nginx/https later

# --- directory to store signed contracts (.docx / .pdf) ---
SIGNED_DIR = BASE_DIR / "signed_docs"
SIGNED_DIR.mkdir(exist_ok=True)

def _clean_money(val) -> int:
    """
    '200 000 руб' → 200000
    176000        → 176000
    """
    if isinstance(val, (int, float)):
        return int(val)
    digits = "".join(ch for ch in str(val) if ch.isdigit())
    return int(digits) if digits else 0

def parse_requisites(raw: str) -> dict:
    """Парсит строку реквизитов и возвращает словарь
    {
        inn, kpp, account, bank, bic, address
    }
    Понимает форматы:
      • ИНН 1234567890  КПП 123456789
      • р/с 40702810…  или р/сч 40702810…  или расч. счёт: 4070…
      • БИК: 044525974  (с двоеточием или без)
      • в банке Тинькофф  |  банк Газпромбанк
      • Юр. адрес: …  |  адрес: …
    """
    if not raw:
        return {}

    res = {}
    rx = lambda p: re.search(p, raw, re.I)

    # ИНН и КПП
    m = rx(r"ИНН\s*([0-9]{6,15})")
    if m:
        res["inn"] = m.group(1)
    m = rx(r"КПП\s*([0-9]{4,10})")
    if m:
        res["kpp"] = m.group(1)

    # расчётный счёт: р/с, р/сч, р с, расч.* счёт
    m = rx(r"(?:р[\s/]?сч?\.?|расч.*сч[её]т)[:№]*\s*([0-9]{15,26})")
    if m:
        res["account"] = m.group(1)

    # Банк – всё после «банк» / «в банке» / «в» до запятой
    m = rx(r"(?:в\s+(?:банке\s+)?|банк\s+)([^,\n]+)")
    if m:
        res["bank"] = m.group(1).strip("«»\"")

    # БИК
    m = rx(r"БИК[:№\s]*([0-9]{6,9})")
    if m:
        res["bic"] = m.group(1)

    # Юридический адрес
    m = rx(r"(?:юр\.?\s*адрес|адрес)[:\s]+(.+)$")
    if m:
        res["address"] = m.group(1).strip()

    return res

# --- PDF act generation ---
def generate_act(order: dict, our_role: str) -> str:
    """
    Собирает и конвертирует договор-заявку в PDF.

    our_role:
        'cust'  – акт для заказчика (мы выступаем Исполнителем)
        'exec'  – акт для исполнителя (мы выступаем Заказчиком)

    Возвращает относительный URL-путь к итоговому PDF.
    """
    # ——— выбираем шаблон ———
    if our_role == "cust":
        tpl_path = ACT_TEMPL_CUST
    else:
        tpl_path = (
            ACT_TEMPL_EXEC
            if ACT_TEMPL_EXEC.exists()
            else next(
                (p for p in ACT_TEMPL_EXEC.parent.glob("template_exec*.docx")),
                ACT_TEMPL_EXEC,
            )
        )

    # ----- выбираем наши реквизиты в зависимости от флага VAT -----
    # vat = True  → работа «с НДС», акт от имени ООО
    # vat = False → «без НДС», акт от имени ИП
    vat_flag = bool(order.get("vat", True))      # default True (ООО)

    if vat_flag:          # с НДС → ООО «ТехноЛогистика»
        our_name = OUR_COMPANY_NAME
        our_req  = OUR_REQUISITES
        our_dir  = OUR_DIRECTOR_NAME
        our_addr = OUR_LEGAL_ADDRESS
    else:                 # без НДС → ИП Хейгетян
        our_name = IP_COMPANY_NAME
        our_req  = IP_REQUISITES
        our_dir  = IP_DIRECTOR_NAME
        our_addr = IP_LEGAL_ADDRESS  # если адрес ИП отличается; иначе можно OUR_LEGAL_ADDRESS

    # ---------- финансовые значения ----------
    price_cust  = order.get("original_amt", 0)
    price_exec  = order.get("final_amt", round(price_cust * 0.88, 2))
    price_total = price_cust if our_role == "cust" else price_exec

    # ---------- helper ----------
    def _normalize_requisites(raw: str) -> str:
        """
        Унифицирует строку реквизитов:
        • добавляет метки «ИНН», «КПП», «р/с», «БИК» там, где их нет;
        • не дублирует их, если они уже присутствуют;
        • оставшуюся «хвостовую» часть (адрес / к/с) добавляет как есть.
        """
        if not raw:
            return ""

        parts = [p.strip() for p in re.split(r"[;,]", raw) if p.strip()]

        inn  = parts[0] if len(parts) > 0 else ""
        kpp  = parts[1] if len(parts) > 1 else ""
        acc  = parts[2] if len(parts) > 2 else ""
        bank = parts[3] if len(parts) > 3 else ""
        bic  = parts[4] if len(parts) > 4 else ""
        tail = ", ".join(parts[5:]) if len(parts) > 5 else ""

        bits = []
        if inn and "ИНН" not in inn.upper():
            bits.append(f"ИНН {inn}")
        elif inn:
            bits.append(inn)

        # --- KPP / OGRNIP ---
        if kpp:
            upper_kpp = kpp.upper()
            if "КПП" in upper_kpp or "ОГРНИП" in upper_kpp:
                # метка уже присутствует (оставляем без изменений)
                bits.append(kpp)
            else:
                # Определяем, что это КФХ/ОГРНИП (длина 15) или обычный КПП
                if len(re.sub(r"\D", "", kpp)) == 15:
                    bits.append(f"ОГРНИП {kpp}")
                else:
                    bits.append(f"КПП {kpp}")

        if acc:
            if "р/с" in acc or "р/c" in acc or "рс" in acc.lower():
                bits.append(acc)
            else:
                bits.append(f"р/с {acc}")

        if bank:
            bank_clean = re.sub(r"[«»]", "", bank.lstrip("в ").strip())
            if bank_clean.lower().startswith(("пao", "ao", "ооо", "банк")):
                bits.append(f"в банке {bank_clean}")
            else:
                bits.append(bank_clean)

        if bic:
            if "бик" in bic.lower():
                bits.append(bic)
            else:
                bits.append(f"БИК {bic}")

        if tail:
            bits.append(tail)

        # ---- final tidy‑up: remove stray «» quotes, duplicate commas, spaces ----
        final = ", ".join(bits)
        final = re.sub(r"[«»“”\"']", "", final)         # убираем любые кавычки
        final = re.sub(r",\s*,", ", ", final)           # двойные запятые
        final = re.sub(r"\s+,", ",", final)             # пробел перед запятой
        final = re.sub(r",\s+$", "", final)             # запятая в конце
        return final.strip()

    # ---------- ensure lists are real Python lists ----------
    def _j2list(val):
        """Return [] for None, decode JSON-string to list, or pass-through list."""
        if val is None:
            return []
        if isinstance(val, str):
            try:
                return json.loads(val)
            except Exception:
                return []
        return val

    cars    = _j2list(order.get("cars"))
    loads   = _j2list(order.get("loads"))
    unloads = _j2list(order.get("unloads"))

    # ---------- контекст для шаблона ----------
    ctx = {
        # общие
        "order_id":          order["id"],
        "sign_date":         datetime.utcnow().strftime("%d.%m.%Y"),
        "vin_list":          ", ".join([c.get("vin","") for c in cars]),
        "price_total":       f"{price_total:,}".replace(",", " "),
        "pay_terms":         order.get("pay_terms", ""),
        "VAT_FLAG":         "с НДС" if order.get("vat") else "без НДС",
        "PAY_TERMS":        order.get("pay_terms", ""),
        # реквизиты и директора
        "cust_requisites":    (order.get("cust_requisites") or "").strip(),
        "cust_director":      (order.get("cust_director")    or "").strip(),
        "carrier_requisites": (order.get("carrier_requisites") or "").strip(),
        "carrier_director":   (order.get("carrier_director")  or order.get("carrier_sign_name") or "").strip(),
        # маршрутные точки
        "loads":   loads,
        "unloads": unloads,
        # динамические подписи сторон
        "cust_company_name": order.get("cust_company_name", ""),
        "carrier_company":   order.get("carrier_company", ""),
        # транспорт / водитель
        "driver_fio":        order.get("driver_fio", ""),
        "driver_passport":   order.get("driver_passport", ""),
        "driver_license":    order.get("driver_license", "—") or "—",
        "truck_reg":         order.get("truck_reg", ""),
        "trailer_reg":       order.get("trailer_reg", ""),
        "truck_model":       order.get("truck_model", ""),
        "trailer_model":     order.get("trailer_model", ""),
        "insurance_policy":  order.get("insurance_policy", ""),
        "truck_info":        order.get("truck_info", ""),
        "trailer_info":      order.get("trailer_info", ""),
        "cust_address":     order.get("cust_address", ""),
        "carrier_address":  order.get("carrier_address", ""),
        "our_address":      OUR_LEGAL_ADDRESS,
    }
        # --- очистим реквизиты от лишних кавычек, дублирующих меток и т.д. ---
    ctx["cust_requisites"]    = _normalize_requisites(ctx.get("cust_requisites", ""))
    ctx["carrier_requisites"] = _normalize_requisites(ctx.get("carrier_requisites", ""))
    # remove any "🏢 Компания…" lines accidentally carried in
    ctx["cust_requisites"]   = re.sub(r"🏢.*", "", ctx.get("cust_requisites", ""))
    ctx["carrier_requisites"] = re.sub(r"🏢.*", "", ctx.get("carrier_requisites", ""))
    # also clean accidental bot messages from director/signature fields
    for key in ("cust_director", "carrier_director", "cust_sign_name", "carrier_sign_name"):
        if "🏢" in ctx.get(key, ""):
            ctx[key] = re.sub(r"🏢.*", "", ctx[key]).strip()
    # --- fallback: если carrier_requisites пуст, берём executor_requisites ---
    if not ctx.get("carrier_requisites"):
        ctx["carrier_requisites"] = ctx.get("executor_requisites", "")
    # ---------- extract individual customer requisites ----------
    raw_cust_req = ctx.get("cust_requisites", "")
    cust_parsed = parse_requisites(raw_cust_req)
    ctx["cust_inn"]       = cust_parsed.get("inn", "")
    ctx["cust_kpp"]       = cust_parsed.get("kpp", "")
    ctx["cust_account"]   = cust_parsed.get("account", "")
    ctx["cust_bank_name"] = cust_parsed.get("bank", "")
    ctx["cust_bic"]       = cust_parsed.get("bic", "")
# адрес заказчика, если нашли
    if cust_parsed.get("address"):
        ctx["cust_address"] = cust_parsed["address"]
# -- если всё равно нет cust_address, достаем из order
    if not ctx.get("cust_address"):
        ctx["cust_address"] = order.get("cust_address", "")
    # --- fallback: ищем ИНН/КПП регуляркой, если split не сработал ---
    if not ctx["cust_inn"]:
        m_inn = re.search(r"ИНН\s*([0-9]{6,15})", raw_cust_req, re.I)
        if m_inn:
            ctx["cust_inn"] = m_inn.group(1)
    if not ctx["cust_kpp"]:
        m_kpp = re.search(r"КПП\s*([0-9]{4,10})", raw_cust_req, re.I)
        if m_kpp:
            ctx["cust_kpp"] = m_kpp.group(1)

    # --- ensure we have customer INN/KPP/addr even if regex failed ---
    if not ctx["cust_inn"]:
        ctx["cust_inn"] = order.get("cust_inn", "")
    if not ctx["cust_kpp"]:
        ctx["cust_kpp"] = order.get("cust_kpp", "")
    if not ctx.get("cust_address"):
        ctx["cust_address"] = order.get("cust_address", "")

    # ---------- extract individual carrier requisites ----------
    raw_car_req = ctx.get("carrier_requisites", "")
    car_parsed = parse_requisites(raw_car_req)
    ctx["carrier_inn"]       = car_parsed.get("inn", ctx.get("carrier_inn",""))
    ctx["carrier_kpp"]       = car_parsed.get("kpp", ctx.get("carrier_kpp",""))
    ctx["carrier_account"]   = car_parsed.get("account", ctx.get("carrier_account",""))
    ctx["carrier_bank_name"] = car_parsed.get("bank", ctx.get("carrier_bank_name",""))
    ctx["carrier_bic"]       = car_parsed.get("bic", ctx.get("carrier_bic",""))
    if car_parsed.get("address"):
        ctx["carrier_address"] = car_parsed["address"]
    if our_role == "cust":          # договор для заказчика, мы Исполнитель
        # очистим левую колонку (данные заказчика) — заполнятся клиентом
        ctx["carrier_company"]    = our_name
        ctx["carrier_requisites"] = our_req
        ctx["carrier_director"]   = our_dir
        ctx["our_address"]        = our_addr
        if not ctx.get("cust_sign_name"):
            ctx["cust_sign_name"] = ctx.get("cust_director", "")
        # ---- split our requisites (carrier side) into separate fields ----
        raw_carrier_req = our_req
        parts = re.split(r"[;,]\s*", raw_carrier_req)
        ctx["carrier_inn"]       = (parts[0] or "").replace("ИНН", "").strip() if len(parts)>0 else ""
        ctx["carrier_kpp"]       = (parts[1] or "").replace("КПП", "").strip() if len(parts)>1 else ""
        ctx["carrier_account"]   = parts[2].strip() if len(parts)>2 else ""
        ctx["carrier_bank_name"] = parts[3].strip() if len(parts)>3 else ""
        ctx["carrier_bic"]       = parts[4].strip() if len(parts)>4 else ""
        ctx["carrier_address"]   = our_addr
        # --- ИП не имеет КПП: очищаем значение, чтобы строка удалилась пост‑процессингом
        if not vat_flag:
            pass  # удалено: не затираем carrier_kpp, обработка ниже
    else:                          # договор для исполнителя, мы Заказчик
        ctx["cust_company_name"]  = our_name
        ctx["cust_requisites"]    = our_req
        ctx["cust_director"]      = our_dir
        ctx["our_address"]        = our_addr
        # ---- split our requisites (customer side) into separate fields ----
        raw_cust_req = our_req
        parts = re.split(r"[;,]\s*", raw_cust_req)
        ctx["cust_inn"]       = (parts[0] or "").replace("ИНН", "").strip() if len(parts)>0 else ""
        ctx["cust_kpp"]       = (parts[1] or "").replace("КПП", "").strip() if len(parts)>1 else ""
        ctx["cust_account"]   = parts[2].strip() if len(parts)>2 else ""
        ctx["cust_bank_name"] = parts[3].strip() if len(parts)>3 else ""
        ctx["cust_bic"]       = parts[4].strip() if len(parts)>4 else ""
        ctx["cust_address"]   = our_addr
        ctx["cust_sign_name"] = our_dir
        ctx["cust_account"]   = parts[2].strip() if len(parts)>2 else ctx.get("cust_account", "")
        ctx["cust_bank_name"] = parts[3].strip() if len(parts)>3 else ctx.get("cust_bank_name", "")
        # --- копируем недостающие реквизиты из executor_* полей ---
        for src, dst in (
            ("executor_inn",        "carrier_inn"),
            ("executor_kpp",        "carrier_kpp"),
            ("executor_bank_rs",    "carrier_account"),
            ("executor_bank_name",  "carrier_bank_name"),
            ("executor_bank_bic",   "carrier_bic"),
            ("executor_address",    "carrier_address"),
        ):
            if not ctx.get(dst):
                ctx[dst] = ctx.get(src, "")
    # --- корректировка ценовых полей для шаблона ---
    if our_role == "cust":            # договор для заказчика (мы Исполнитель)
        ctx["original_amt"] = f"{price_cust:,}".replace(",", " ")
        ctx["final_amt"]    = ctx["original_amt"] + " руб."
        ctx["our_role_label"] = "Исполнитель"
    else:                             # договор для исполнителя (мы Заказчик)
        ctx["original_amt"] = f"{price_exec:,}".replace(",", " ")
        ctx["final_amt"]    = ctx["original_amt"] + " руб."
        ctx["our_role_label"] = "Заказчик"

    # Remove unintended override of cust_sign_name (do not set here)
    # ---------- генерация через docxtpl ----------
    doc = DocxTemplate(tpl_path)
    doc.render(ctx)

    final_docx = ACTS_DIR / f"act_{order['id']}_{our_role}.docx"
    # ------------------------------------------------------------------
    # Post‑processing: remove КПП row for IP or when value is blank
    # ------------------------------------------------------------------
    def _drop_kpp_row(document, is_ip: bool):
        """
        • Если value пусто («», «—») → строку удаляем всегда.
        • Для ИП: не удаляем по is_ip, только если пусто.
        """
        for tbl in document.tables:
            for row in tbl.rows:
                cells = row.cells
                if len(cells) < 2:
                    continue
                label = cells[0].text.strip().upper()
                value = cells[1].text.strip()
                value_clean = value.replace("—", "").replace("-", "").strip()
                # обрабатываем только строки «КПП …»
                if not re.search(r"КПП", label, re.I):
                    continue
                # удаляем, если значение пустое / «—»
                if not value_clean:
                    tr = row._tr
                    tr.getparent().remove(tr)

    _drop_kpp_row(doc, is_ip=not vat_flag)
    # ------------------------------------------------------------------
    # Remove textual fragment "КПП —" (for IP where КПП is not applicable)
    # ------------------------------------------------------------------
    def _remove_kpp_dash(document):
        """Стирает «КПП —» (с возможной запятой) из всех параграфов."""
        pattern = re.compile(r"КПП\s+[—\-–−](?:,\s*)?", re.I)
        for p in document.paragraphs:
            if "КПП" in p.text:
                p.text = pattern.sub("", p.text).strip()

        # таблицы тоже могут иметь параграфы
        for tbl in document.tables:
            for row in tbl.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        if "КПП" in p.text:
                            p.text = pattern.sub("", p.text).strip()

    _remove_kpp_dash(doc)
    # ------------------------------------------------------------------
    # ------------------------------------------------------------------
    # Replace label «КПП» → «ОГРНИП» in carrier table for IP variant (robust: match КПП anywhere in cell)
    # ------------------------------------------------------------------
    if not vat_flag:
        for tbl in doc.tables:
            for row in tbl.rows:
                # ищем «КПП» хоть где‑то в первой ячейке, даже если перед ним перевод строки
                if "КПП" in row.cells[0].text:
                    row.cells[0].text = re.sub(r"КПП", "ОГРНИП", row.cells[0].text, 1)

    doc.save(final_docx)
    return str(final_docx.relative_to(BASE_DIR))


# ——————————————————— helpers ———————————————————
def _replace_tags(document: Document, mapping: dict) -> None:
    """Глобальная замена {{tag}} → value во всём документе."""
    for paragraph in document.paragraphs:
        for tag, value in mapping.items():
            if tag in paragraph.text:
                paragraph.text = paragraph.text.replace(tag, str(value))

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for tag, value in mapping.items():
                        if tag in p.text:
                            p.text = p.text.replace(tag, str(value))


def _route_map(order: dict) -> dict:
    """Теги load*/unload* для первых двух точек маршрута."""
    def _json2list(x):
        if x is None:
            return []
        if isinstance(x, str):
            try:
                return json.loads(x)
            except Exception:
                return []
        return x

    loads   = _json2list(order.get("loads"))
    unloads = _json2list(order.get("unloads"))
    res: dict[str, str] = {}

    for idx, l in enumerate(loads[:2], 1):
        res[f"{{{{load{idx}_point}}}}"]   = l.get("place", "")
        res[f"{{{{load{idx}_date}}}}"]    = l.get("date", "")
        vins = l.get("vins") or []
        res[f"{{{{load{idx}_vins}}}}"]    = ", ".join(vins) if isinstance(vins, list) else vins
        res[f"{{{{load{idx}_contact}}}}"] = l.get("contact", "")

    for idx, u in enumerate(unloads[:2], 1):
        res[f"{{{{unload{idx}_point}}}}"] = u.get("place", "")
        res[f"{{{{unload{idx}_date}}}}"]  = u.get("date", "")
        vins = u.get("vins") or []
        res[f"{{{{unload{idx}_vins}}}}"]  = ", ".join(vins) if isinstance(vins, list) else vins
        res[f"{{{{unload{idx}_contact}}}}"] = u.get("contact", "")
    return res


def _cars_table(cars) -> str:
    """Формирует markdown‑таблицу из списка автомобилей."""
    if not cars:
        return ""
    if isinstance(cars, str):
        try:
            cars = json.loads(cars)
        except Exception:
            return ""
    lines = ["Марка | VIN", "----- | ----"]
    for c in cars:
        lines.append(f"{c.get('brand','')} {c.get('model','')} | {c.get('vin','')}")
    return "\n".join(lines)
# ---------- route parser ----------
_route_re = re.compile(r"^\s*(.+?)\s*[—\-]\s+(.+?)(?:,|$)", re.U)

def extract_route(text: str) -> tuple[str, str] | None:
    """
    Попытаться вытащить «ГородA — ГородB» или «ГородA-ГородB» из начала строки.
    Возвращает (from_city, to_city) либо None.
    """
    m = _route_re.match(text)
    if not m:
        return None
    from_city = m.group(1).strip().title()
    to_city   = m.group(2).strip().title()
    return from_city, to_city

from sheets import add_request_row as add_record, sheet
from database import database, agents, orders
from database import companies as company


from sqlalchemy.dialects.postgresql import insert  # for UPSERT

# ---------- financial constants ----------
COMPANY_MARKUP  = 0.12   # 12 % общая надбавка к цене исполнителя
CUST_SHARE      = 0.05   # 5 % от стоимости — агент‑заказчик
EXEC_SHARE      = 0.03   # 3 % от стоимости — агент‑исполнитель
PLATFORM_SHARE  = 0.04   # 4 % от стоимости — платформе
# ---------- our company ----------
OUR_COMPANY_NAME   = "ООО «ТехноЛогистика»"
OUR_REQUISITES     = "7700000000, 770001001, 40702810900000000000, АО «Банк», 044525000"
OUR_DIRECTOR_NAME  = "Новиков Е.О."
OUR_DIRECTOR_POS   = "Генеральный директор"
OUR_LEGAL_ADDRESS = "г. Москва, вн.тер.г. Муниципальный округ Красносельский, пер. Уланский, д. 22, стр 1, помещ. 41Н/6"
# --- our Individual Entrepreneur variant (без НДС) ---
# --- Individual Entrepreneur variant (без НДС) ---
IP_COMPANY_NAME  = "ИП Хейгетян Е.В."
IP_REQUISITES = (
    "ИНН 616610295207, ОГРНИП 325619600016460, "
    "40802810220000542247, ООО «Банк Точка», "
    "044525104, 30101810745374525104"
)
IP_DIRECTOR_NAME = "Хейгетян Е.В."
IP_LEGAL_ADDRESS = "344029, г. Ростов-на-Дону, ул. Металлургическая 15"
# ---------- statuses ----------
STATUSES_ALL      = ("active", "confirmed", "in_progress", "done", "cancelled", "paid")
CURRENT_STATUSES  = ("active", "confirmed", "in_progress", "done", "cancelled")  # всё, кроме paid
ARCHIVE_STATUSES  = ("paid",)      # в архив попадают только оплаченные
# ---------- Persistence helpers --------------------------
BASE_DIR = "/home/gpt/technologist_crm"
AGENTS_FILE = os.path.join(BASE_DIR, "agents.json")

def load_agents() -> dict[int, dict]:
    if os.path.exists(AGENTS_FILE):
        with open(AGENTS_FILE, "r", encoding="utf-8") as f:
            raw = json.load(f)
            return {int(k): v for k, v in raw.items()}
    return {}

def save_agents() -> None:
    with open(AGENTS_FILE, "w", encoding="utf-8") as f:
        json.dump(agents_db, f, ensure_ascii=False, indent=2)
# ----------------------------------------------------------

# ---------- INVITES: single-use invite token mechanism ----------
INVITES_FILE = os.path.join(BASE_DIR, "invites.json")

def load_invites() -> dict[str, dict]:
    if os.path.exists(INVITES_FILE):
        with open(INVITES_FILE, "r", encoding="utf-8") as f:
            return json.load(f)
    return {}

def save_invites() -> None:
    with open(INVITES_FILE, "w", encoding="utf-8") as f:
        json.dump(invites_db, f, ensure_ascii=False, indent=2)

invites_db: dict[str, dict] = load_invites()

# ---------- In‑memory “DB” (simple demo, replace by real DB) ----
agents_db: dict[int, dict] = load_agents()
orders_db: dict[int, dict] = {}
# ----------------------------------------------------------------

# -------------------- Telegram notify helper --------------------

BOT_TOKEN = "7626649459:AAFYfJrC31GzZgEKNQUbhf11wbP8dN5mhgU"
# TG‑group (private) where every signed contract is forwarded
SIGNED_GROUP_ID = int(os.getenv("SIGNED_GROUP_ID", "-4864154859"))  # chat_id группы «CRM договора»

def run_telegram_bot():
    application = ApplicationBuilder().token(BOT_TOKEN).build()

    # стартовый хендлер /start
    async def start(update: Update, context: ContextTypes.DEFAULT_TYPE):
        await update.message.reply_text("Привет! Я бот CRM.")

    start_handler = CommandHandler("start", start)
    application.add_handler(start_handler)

    # регистрируем конверсацию «Закрыть заявку»
    register_close_conversation(application, start_handler)

    application.run_polling()

# запускаем Telegram‑бота в отдельном потоке, чтобы uvicorn не блокировался
threading.Thread(target=run_telegram_bot, daemon=True).start()
print("🤖 Бот запущен и ждёт сообщений...")
BOT_USERNAME = "technologist_nav_bot"  # ← replace with your bot’s username

def send_notification_to_executor(
    tg_id: int,
    text: str,
    order_id: int | None = None,
    reward: int | None = None,
) -> None:
    """
    Отправить исполнителю (или заказчику) уведомление.

    • Если передан reward — добавляем строку «🏆 Ваш бонус: +XXXX ₽».
    • Если передан order_id — добавляем inline‑кнопку «Закрыть».
    """
    full_text = text
    if reward is not None:
        full_text += f"\n🏆 Ваш бонус: +{reward} ₽"

    url = f"https://api.telegram.org/bot{BOT_TOKEN}/sendMessage"
    payload: dict = {"chat_id": tg_id, "text": full_text}

    if order_id is not None:
        payload["reply_markup"] = {
            "inline_keyboard": [[
                {"text": "Закрыть", "callback_data": f"close_{order_id}"}
            ]]
        }

    try:
        requests.post(url, json=payload, timeout=4)
    except Exception as e:
        print("Telegram send error:", e)
# ----------------------------------------------------------------

# --- helper: post signed contract to archive group ---
def _post_to_signed_group(file_path: Path, caption: str):
    """Forward a signed contract to the 📑 Подписанные договоры group."""
    try:
        url = f"https://api.telegram.org/bot{BOT_TOKEN}/sendDocument"
        with open(file_path, "rb") as f:
            requests.post(
                url,
                data={"chat_id": SIGNED_GROUP_ID, "caption": caption},
                files={"document": f},
                timeout=10,
            )
    except Exception as e:
        print("Signed group post error:", e)

app = FastAPI()
 
# --- CORS  ----------------------------------------------------------
app.add_middleware(
    CORSMiddleware,
    allow_origins=["http://localhost:5173"],  # адрес фронта (React/Vite)
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)
# --------------------------------------------------------------------

# Serve generated PDF acts
app.mount("/acts", StaticFiles(directory=ACTS_DIR), name="acts")

@app.on_event("startup")
async def _startup():
    await database.connect()
    # --- DEMO: создаём 3 тестовые заявки через /admin/mock_order для telegram_id=7823236991 ---
    # --- DEMO: создаём 3 тестовые заявки через /admin/mock_order для telegram_id=6835069941 ---
    # (удалить или закомментировать блок ниже на проде)
    # try:
    #     for _ in range(3):
    #         await admin_mock_order(telegram_id=6835069941)
    # except Exception as e:
    #     print("Mock order creation error:", e)

@app.on_event("shutdown")
async def _shutdown():
    await database.disconnect()

# ---------------------------- MODELS ----------------------------

# --------- COMPANY MODEL ----------
class CompanyIn(BaseModel):
    inn: str
    kpp: str | None = ""
    name: str | None = ""
    director: str | None = ""
    bank: str | None = ""
    address: str | None = ""

class AgentRegistration(BaseModel):
    telegram_id: int
    name: str
    agent_type: str  # 'заказчик' | 'исполнитель'
    phone: str | None = ""
# --------- ORDER MODEL ----------
class Order(BaseModel):
    telegram_id: int
    message: str
    original_amt: Optional[int] = None
    cust_requisites: Optional[str] = None
    cars:    Optional[list] = None
    loads:   Optional[list] = None
    unloads: Optional[list] = None
    pay_terms: Optional[str] = None   # форма и срок оплаты
    insurance_policy: Optional[str] = None  # № страхового полиса
    driver_license:  Optional[str] = None   # номер ВУ
    truck_model:     Optional[str] = None   # марка тягача
    trailer_model:   Optional[str] = None   # марка прицепа
    vat: Optional[bool] = True   # True = с НДС (ООО), False = без НДС (ИП)
    

    # --- CUSTOMER‑specific fields ---
    cust_company_name: Optional[str] = None      # название ООО заказчика (карточка предприятия)
    cust_director:     Optional[str] = None      # ФИО директора заказчика

    # --- EXECUTOR‑specific fields (заполняет агент‑исполнитель при закрытии) ---
    driver_fio:        Optional[str] = None      # ФИО водителя
    driver_passport:   Optional[str] = None      # серия/номер паспорта водителя
    truck_reg:         Optional[str] = None      # гос‑номер тягача
    trailer_reg:       Optional[str] = None      # гос‑номер прицепа
    carrier_requisites: Optional[str] = None     # ИНН/КПП/банк исполнителя

class CloseOrderRequest(BaseModel):
    order_id: int
    executor_id: int
    driver_fio:        Optional[str] = None
    driver_passport:   Optional[str] = None   # серия / номер паспорта водителя
    truck_reg:         Optional[str] = None   # гос‑номер тягача
    trailer_reg:       Optional[str] = None   # гос‑номер прицепа
    carrier_requisites: Optional[str] = None  # ИНН/КПП/банк
    insurance_policy: Optional[str] = None   # ОСАГО / CMR-страховка
    driver_license:  Optional[str] = None
    truck_model:     Optional[str] = None
    trailer_model:   Optional[str] = None
    carrier_company:   Optional[str] = None      # название компании‑перевозчика
    carrier_director:  Optional[str] = None      # ФИО директора перевозчика
    truck_info:       Optional[str] = None      # марка + гос‑номер тягача (для шаблона)
    trailer_info:     Optional[str] = None      # марка + гос‑номер прицепа (для шаблона)
    # новые поля
    loads:   Optional[list] = None            # [{"place": "...", "date": "..."}]
    unloads: Optional[list] = None
    vat: Optional[bool] = None   # True = с НДС (ООО), False = без НДС (ИП)

# --------- INVITE TOKEN MODELS ----------
# --------- INVITE TOKEN MODELS ----------
class InviteCreate(BaseModel):
    role: str  # заказчик | исполнитель
    ttl_hours: int = 72

class InviteClaim(BaseModel):
    telegram_id: int
    token: str


# ---------- ADD ORDER endpoint -------------------------------------------------
@app.post("/add_order")
async def add_order(order: Order):
    """
    Заказчик публикует заявку.
    Сохраняем в БД, пишем в Google‑Sheets и уведомляем всех исполнителей.
    """
    original_msg = order.message.strip()

    # --- 1. Извлекаем маршрут ---
    route = extract_route(original_msg)
    if not route:
        raise HTTPException(
            status_code=400,
            detail="Укажите маршрут в формате «ГородA — ГородB» перед суммой."
        )
    route_from, route_to = route

    # --- 2. Сумма ---
    if order.original_amt is not None:
        original_amt = order.original_amt
    else:
        # Ищем все числовые блоки, допускаем обычный пробел, NBSP (U+00A0) и тонкий NBSP (U+202F)
        price_candidates = re.findall(r"\d[\d\s\u00A0\u202F]{2,12}", original_msg)
        nums = [_clean_money(p) for p in price_candidates]
        original_amt = max(nums) if nums else 0

        if original_amt == 0:
            raise HTTPException(
                status_code=400,
                detail="Укажите original_amt или впишите сумму в текст, например «… 45000 руб»"
            )

    # --- 3. Расчёт цен и вознаграждений ---
    margin        = int(original_amt * COMPANY_MARKUP)
    reward_cust   = int(original_amt * CUST_SHARE)
    reward_exec   = int(original_amt * EXEC_SHARE)
    fee_platform  = int(original_amt * PLATFORM_SHARE)
    final_amt     = original_amt - margin

    # --- 4. Сохраняем в Postgres ---
    pg_query = orders.insert().values(
        telegram_id   = order.telegram_id,
        message       = original_msg,
        original_amt  = original_amt,
        final_amt     = final_amt,
        reward_cust   = reward_cust,
        reward_exec   = reward_exec,
        fee_platform  = fee_platform,
        status        = "active",
        created_at    = datetime.utcnow(),
        cust_requisites    = order.cust_requisites,
        cust_company_name  = order.cust_company_name,
        cust_director      = order.cust_director,
        cars    = json.dumps(order.cars,    ensure_ascii=False) if order.cars    else None,
        loads   = json.dumps(order.loads,   ensure_ascii=False) if order.loads   else None,
        unloads = json.dumps(order.unloads, ensure_ascii=False) if order.unloads else None,
        pay_terms         = order.pay_terms,
        insurance_policy  = order.insurance_policy,
        driver_license    = order.driver_license,
        truck_model       = order.truck_model,
        trailer_model     = order.trailer_model,
        vat              = order.vat,
    )
    order_id = await database.execute(pg_query)

    # --- 5. Запись в Google‑Sheets (best-effort) ---
    try:
        add_record(
            agent_type   = "заказчик",
            name         = agents_db.get(order.telegram_id, {}).get("name", "—"),
            tg_id        = order.telegram_id,
            message      = f"[{order_id}] {original_msg}",
            original_amt = original_amt,
            final_amt    = final_amt,
            status       = ""
        )
    except Exception as e:
        print("Sheets add_record error:", e)

    # --- 6. Формируем уведомление для исполнителей ----------
    # Сначала общий заголовок с ID и маршрутом
    route_txt = f"{route_from} — {route_to}"
    header = f"[{order_id}] {route_txt}"

    # --- build structured message: loads, unloads, cargo, price, payment ---
    sections = [header]

    # Пункты погрузки
    if order.loads:
        sections.append("📍 Пункты погрузки:")
        for i, item in enumerate(order.loads, start=1):
            sections.append(f"  {i}. {item.get('place','')} ({item.get('date','')})")
        sections.append("")  # blank line

    # Пункты выгрузки
    if order.unloads:
        sections.append("📍 Пункты выгрузки:")
        for i, item in enumerate(order.unloads, start=1):
            sections.append(f"  {i}. {item.get('place','')} ({item.get('date','')})")
        sections.append("")  # blank line

    # Описание груза
    if order.cars:
        first = order.cars[0] if order.cars else {}
        brand = (first.get("brand") or "").strip()
        model = (first.get("model") or "").strip()
        if brand or model:
            brand_model = f"{brand} {model}".strip()
            cargo_txt = f"{len(order.cars)} × {brand_model}"
        else:
            m = re.search(r",\s*([^,]+?),\s*(\d+)\s+авто", original_msg, re.I)
            if m:
                brand_model = m.group(1).strip()
                qty = m.group(2)
                cargo_txt = f"{qty} × {brand_model}"
            else:
                m2 = re.search(r"(\d+)\s+авто", original_msg, re.I)
                cargo_txt = f"{m2.group(1)} авто" if m2 else ""
        sections.append(f"🚚 Груз: {cargo_txt}")

    # Цена исполнителю (без НДС)
    exec_price_text = f"{final_amt:,} руб".replace(",", " ")
    sections.append(f"💵 Цена (для исполнителя): {exec_price_text}")

    # Условия оплаты
    if order.pay_terms:
        sections.append(f"💳 Условия оплаты: {order.pay_terms}")

    notify_text = "\n".join(sections)

    # бонус = 3 % от ИСХОДНОЙ цены (до вычета 12 %)
    bonus_exec = int(round(original_amt * 0.03))

    # --- 7. Рассылаем исполнителям ---
    exec_rows = await database.fetch_all(
        agents.select().where(agents.c.agent_type == "исполнитель")
    )
    for r in exec_rows:
        tg_id = r["telegram_id"]
        # одно сообщение; send_notification… сам добавит строку бонуса
        send_notification_to_executor(
            tg_id,
            notify_text,
            order_id=order_id,          # для кнопки «Закрыть»
            reward=bonus_exec,          # строка бонуса
        )

    # --- 8. Ответ заказчику ---
    return {
        "status": "Заявка принята",
        "message": f"Ваш номер заявки: {order_id}"
    }
# ---------------------------------------------------------------------------

@app.post("/register_agent")
async def register_agent(registration: AgentRegistration):
    if registration.agent_type not in ("заказчик", "исполнитель"):
        raise HTTPException(status_code=400, detail="Invalid agent type")

    stmt = (
        insert(agents)
        .values(
            telegram_id   = registration.telegram_id,
            name          = registration.name,
            agent_type    = registration.agent_type,
            phone         = registration.phone or "",
            registered_at = datetime.utcnow(),
        )
        .on_conflict_do_update(
            index_elements=[agents.c.telegram_id],
            set_={
                "name":        registration.name,
                "agent_type":  registration.agent_type,
                "phone":       registration.phone or "",
                "registered_at": agents.c.registered_at  # keep original date
            }
        )
        .returning(agents.c.id)
    )
    await database.execute(stmt)

    # обновляем локальный кэш
    agents_db[registration.telegram_id] = {
        "name": registration.name,
        "agent_type": registration.agent_type,
        "phone": registration.phone or ""
    }
    # persist to disk so the agent survives restarts
    save_agents()

    return {"status": "ok"}


    


@app.post("/close_order")
async def close_order(request: CloseOrderRequest):
    """Исполнитель закрывает заявку. Отмечаем локально и добавляем строку «Закрыта» в таблицу."""
    # ---- VALIDATION: only agent‑исполнитель can close ----
    exec_profile = agents_db.get(request.executor_id)

    # если нет в кэше – пробуем достать из БД
    if exec_profile is None:
        db_row = await database.fetch_one(
            agents.select().where(agents.c.telegram_id == request.executor_id)
        )
        if db_row:
            exec_profile = dict(db_row)
            agents_db[request.executor_id] = {
                "name": exec_profile["name"],
                "agent_type": exec_profile["agent_type"],
                "phone": exec_profile.get("phone", "")
            }

    agent_type = (exec_profile.get("agent_type") or "").strip().lower() if exec_profile else ""
    if agent_type != "исполнитель":
        raise HTTPException(status_code=403, detail="Закрывать заявку может только агент‑исполнитель")

    # --- получаем заявку из кэша или вытаскиваем из Postgres ---
    order = orders_db.get(request.order_id)
    if order is None:
        db_row = await database.fetch_one(
            orders.select().where(orders.c.id == request.order_id)
        )
        if db_row is None:
            raise HTTPException(status_code=404, detail="Order not found")
        order = dict(db_row)
        orders_db[request.order_id] = order

    # --- Блокировка: если заявка уже не "active" или уже занята другим ---
    if order["status"] != "active":
        raise HTTPException(status_code=409, detail="Заявка уже взята или закрыта другим исполнителем")
    if order.get("executor_id") and order.get("executor_id") != request.executor_id:
        raise HTTPException(status_code=409, detail="Заявка уже в работе у другого исполнителя")

    # гарантируем наличие поля id для генерации договора
    order["id"] = request.order_id

    # --- запрет заказчику закрывать собственную заявку ---
    if order["telegram_id"] == request.executor_id:
        raise HTTPException(status_code=403, detail="Заказчик не может закрыть собственную заявку")

    # локальная отметка
    order["status"] = "confirmed"
    order["executor_id"] = request.executor_id
    order["carrier_requisites"] = request.carrier_requisites
    order["driver_passport"] = request.driver_passport
    order["truck_reg"]       = request.truck_reg
    order["trailer_reg"]     = request.trailer_reg
    if request.loads is not None:
        order["loads"] = request.loads

    if request.unloads is not None:
        order["unloads"] = request.unloads
    order["insurance_policy"] = request.insurance_policy
    order["driver_license"]   = request.driver_license
    order["truck_model"]      = request.truck_model
    order["trailer_model"]    = request.trailer_model
    if request.vat is not None:
        order["vat"] = request.vat
    order["carrier_company"]   = request.carrier_company
    order["carrier_sign_name"] = request.carrier_director
    order["driver_fio"]     = request.driver_fio
    order["truck_info"]     = request.truck_info
    order["trailer_info"]   = request.trailer_info

    # --- обновляем статус в Postgres ---
    await database.execute(
        orders.update()
              .where(orders.c.id == request.order_id)
              .values(
                  status="confirmed",
                  executor_id=request.executor_id,
                  driver_fio        = request.driver_fio,
                  driver_passport   = request.driver_passport,
                  truck_reg         = request.truck_reg,
                  trailer_reg       = request.trailer_reg,
                  carrier_requisites= request.carrier_requisites,
                  insurance_policy  = request.insurance_policy,
                  driver_license    = request.driver_license,
                  truck_model       = request.truck_model,
                  trailer_model     = request.trailer_model,
                  loads   = json.dumps(request.loads, ensure_ascii=False) if request.loads is not None else None,
                  unloads = json.dumps(request.unloads, ensure_ascii=False) if request.unloads is not None else None,
                  vat             = request.vat if request.vat is not None else orders.c.vat,
                  updated_at=datetime.utcnow(),
              )
    )

    executor_info = agents_db.get(
        request.executor_id,
        {"name": "—", "agent_type": "исполнитель"}
    )

    # добавляем строку‑лог в таблицу
    try:
        sheet.append_row([
            datetime.now().strftime('%Y-%m-%d %H:%M:%S'),   # Дата/время
            executor_info["agent_type"],                    # тип
            executor_info["name"],                          # имя
            request.executor_id,                            # TG‑ID
            f"[{request.order_id}] {order['message']}",     # текст заявки
            "",                                             # Оригинал
            "",                                             # Итог
            "Закрыта",                                      # Статус
            request.driver_fio or "-"                       # ФИО водителя
        ])
        print(f"✅ Added close‑row for order {request.order_id}")
    except Exception as e:
        print("Sheets append_row error:", e)

    # ---------- уведомляем исполнителя ----------
    try:
        send_notification_to_executor(
            request.executor_id,
            f"✅ Данные приняты!\n"
            f"Заявка №{request.order_id} перешла в статус «Подтверждена». "
            f"Ожидайте проверки документов."
        )
    except Exception as e:
        print("Telegram notify executor error:", e)

    # ---------- уведомляем заказчика ----------
    customer_id = order["telegram_id"]
    try:
        send_notification_to_executor(
            customer_id,
            f"🚚 Исполнитель найден!\n"
            f"Заявка №{request.order_id} перешла в статус «Подтверждена»."
        )
    except Exception as e:
        print("Telegram notify customer error:", e)

    # ---------- генерируем два договора и рассылаем ссылки ----------
    act_for_customer_rel  = None  # договор, где мы — Исполнитель (для Заказчика)
    act_for_executor_rel  = None  # договор, где мы — Заказчик   (для Исполнителя)
    link_cust = None  # ссылка для заказчика
    link_exec = None  # ссылка для исполнителя
    try:
        print("DBG: start generate_act for order", order["id"])
        act_for_customer_rel = generate_act(order, our_role="cust")  # мы Исполнитель
        print("DBG: cust OK ->", act_for_customer_rel)
        act_for_executor_rel = generate_act(order, our_role="exec")  # мы Заказчик
        print("DBG: exec OK ->", act_for_executor_rel)

        link_cust = f"{BASE_URL}/acts/{Path(act_for_customer_rel).name}"
        link_exec = f"{BASE_URL}/acts/{Path(act_for_executor_rel).name}"

        # --- отправляем ссылки пользователям ---
        send_notification_to_executor(customer_id, f"📄 Договор для подписания (мы — Исполнитель):\n{link_cust}")
        send_notification_to_executor(request.executor_id, f"📄 Договор для подписания (мы — Заказчик):\n{link_exec}")
        print("DBG: links sent to users")
    except Exception as e:
        print("Generate/Send act error:", e)

    # ---------- финальный ответ боту ----------
    return {
        "status": "ok",
        "customer_tg": customer_id,
        "executor_tg": request.executor_id,
        "cust_path": act_for_customer_rel,
        "exec_path": act_for_executor_rel,
        "link_cust": link_cust,
        "link_exec": link_exec,
    }

# ---------- Upload signed contract (.docx / .pdf) ----------
@app.post("/orders/{order_id}/signed_document")
async def upload_signed_document(
    order_id: int,
    role: str = "navigator",               # navigator | driver
    filename: str | None = None,
    file: UploadFile = File(...),
):
    """
    Навигатор / Драйвер загружает подписанный договор.
    • Сохраняем файл в SIGNED_DIR
    • Обновляем orders.*_path
    • Если оба файла получены ⇒ status='in_progress'
    • Отправляем документ в группу SIGNED_GROUP_ID
    """
    # --- ensure order exists ---
    row = await database.fetch_one(orders.select().where(orders.c.id == order_id))
    if not row:
        raise HTTPException(status_code=404, detail="Order not found")
    row = dict(row)

    # --- save file ---
    suffix = Path(file.filename or filename or "contract.docx").suffix.lower()
    save_path = SIGNED_DIR / f"{order_id}_{role}{suffix}"
    with open(save_path, "wb") as f:
        f.write(await file.read())

    # --- prepare DB update ---
    values: dict = {}
    role = role.lower()
    if role in ("driver", "exec", "executor"):
        values["signed_exec_path"] = str(save_path.relative_to(BASE_DIR))
    else:
        values["signed_cust_path"] = str(save_path.relative_to(BASE_DIR))

    # check completion
    signed_exec = row.get("signed_exec_path") or values.get("signed_exec_path")
    signed_cust = row.get("signed_cust_path") or values.get("signed_cust_path")
    complete = bool(signed_exec and signed_cust)

    if complete:
        values.update(
            {
                "is_signed_complete": True,
                "signed_uploaded_at": datetime.utcnow(),
                "status": "in_progress",
            }
        )

    # persist
    await database.execute(
        orders.update().where(orders.c.id == order_id).values(**values)
    )

    # archive to group
    _post_to_signed_group(save_path, f"Заявка #{order_id} • подписал {role}")

    # notify parties if both signatures collected
    if complete:
        try:
            send_notification_to_executor(
                row["telegram_id"],
                f"🎉 Оба подписанных договора по заявке №{order_id} получены.\n"
                f"Заявка переведена в статус «В работе».",
            )
            if row.get("executor_id"):
                send_notification_to_executor(
                    row["executor_id"],
                    f"🎉 Оба подписанных договора по заявке №{order_id} получены.\n"
                    f"Заявка переведена в статус «В работе».",
                )
        except Exception as e:
            print("Notify parties error:", e)

    return {"status": "in_progress" if complete else "partial"}

class StatusUpdate(BaseModel):
    status: str  # confirmed | in_progress | done | paid


# ---------- New helper endpoints for bot UI ----------

@app.patch("/orders/{order_id}/status")
async def update_order_status(order_id: int, upd: StatusUpdate):
    """Изменить статус заявки + уведомить заказчика и исполнителя."""
    if upd.status not in STATUSES_ALL:
        raise HTTPException(status_code=400, detail="Invalid status")

    # --- обновляем в БД ---
    await database.execute(
        orders.update()
              .where(orders.c.id == order_id)
              .values(status=upd.status, updated_at=datetime.utcnow())
    )

    # --- получаем свежую строку ---
    row = await database.fetch_one(
        orders.select().where(orders.c.id == order_id)
    )
    if not row:
        raise HTTPException(status_code=404, detail="Order not found")

    row = dict(row)  # convert Record to plain dict

    # если в БД ещё нет executor_id (старая запись) — берём из кэша
    if not row.get("executor_id"):
        cached = orders_db.get(order_id, {})
        row["executor_id"] = cached.get("executor_id")

    # --- оповещаем стороны ---
    STATUS_RU = {
        "active":      "Активна",
        "confirmed":   "Подтверждена",
        "in_progress": "В работе",
        "done":        "Исполнена",
        "cancelled":   "Отменена",
        "paid":        "Оплачена",
    }
    status_txt = STATUS_RU.get(upd.status, upd.status)
    note_text  = f"🔔 Статус вашей заявки №{order_id} изменён на «{status_txt}»"

    # заказчик
    send_notification_to_executor(row["telegram_id"], note_text)

    # исполнитель (если уже указан)
    if row.get("executor_id"):
        send_notification_to_executor(row["executor_id"], note_text)

    # --- акт и выплаты ---
    if upd.status == "confirmed":
        # --- два договора‑заявки ---
        # 1) act_cust_rel → договор для заказчика (мы Исполнитель, полная цена)
        # 2) act_exec_rel → договор для исполнителя (мы Заказчик, цена −12 %)
        act_cust_rel = generate_act(row, our_role="cust")   # для заказчика
        act_exec_rel = generate_act(row, our_role="exec")   # для исполнителя

        link_cust = f"{BASE_URL}/acts/{Path(act_cust_rel).name}"
        link_exec = f"{BASE_URL}/acts/{Path(act_exec_rel).name}"

        # — отправляем заказчику —
        send_notification_to_executor(
            row["telegram_id"],
            f"📄 Договор для подписания (мы — Исполнитель):\n{link_cust}"
        )
        # — отправляем исполнителю —
        if row.get("executor_id"):
            send_notification_to_executor(
                row["executor_id"],
                f"📄 Договор для подписания (мы — Заказчик):\n{link_exec}"
            )
    return {"status": "ok"}

# ---------- ADMIN: change order status via /admin/order/{id}/status/{action} ----------
@app.patch("/admin/order/{order_id}/status/{action}")
async def admin_change_status(order_id: int, action: str):
    """
    Администратор меняет статус заявки любым из допустимых значений.
    Прокси к /orders/{order_id}/status с тем же набором проверок.
    """
    upd = StatusUpdate(status=action)
    return await update_order_status(order_id=order_id, upd=upd)

@app.get("/open_orders")
async def get_open_orders(
    limit: int = 15,
    origin: str | None = None,
    dest: str | None = None,
    cargo_kw: str | None = None,
    min_reward: int | None = None,
):
    """
    Список «Активных» заявок с возможностью фильтра.

    • origin / dest  — строки, ищутся внутри message (без учёта регистра)
    • cargo_kw       — ключевое слово по описанию груза
    • min_reward     — минимальный бонус исполнителю (reward_exec)
    """
    q = (
        orders.select()
              .where(orders.c.status == "active")
              .where(orders.c.executor_id.is_(None))   # показываем только незанятые
              .order_by(orders.c.id.desc())
    )

    if origin:
        if hasattr(orders.c, "route_from"):
            q = q.where(orders.c.route_from.ilike(f"%{origin}%"))
        else:
            q = q.where(orders.c.message.ilike(f"%{origin}%"))

    if dest:
        if hasattr(orders.c, "route_to"):
            q = q.where(orders.c.route_to.ilike(f"%{dest}%"))
        else:
            q = q.where(orders.c.message.ilike(f"%{dest}%"))

    if cargo_kw:
        q = q.where(orders.c.message.ilike(f"%{cargo_kw}%"))
    if min_reward is not None:
        q = q.where(orders.c.reward_exec >= min_reward)

    rows = await database.fetch_all(q.limit(limit))
    # Convert each row (Record) to plain dict first to avoid attribute errors
    return [
        {
            "id":          d["id"],
            "message":     d["message"],
            "final_amt":   d["final_amt"],
            "created_at":  d["created_at"],
            "status":      d["status"],
            "executor_id": d["executor_id"],
        }
        for d in map(dict, rows)
    ]



@app.get("/order/{order_id}")
async def get_order(order_id: int):
    """
    Вернуть детальную информацию по одной заявке (открытой ИЛИ закрытой).

    • Для заказчика: покажет исходное сообщение, итоговую цену (для исполнителя)
      и его вознаграждение.
    • Для исполнителя: дополнительно вернёт driver_fio и его бонус.

    Если запись не найдена — 404.
    """
    row = await database.fetch_one(
        orders.select().where(orders.c.id == order_id)
    )
    if not row:
        raise HTTPException(status_code=404, detail="Order not found")
    d = dict(row)  # convert Record to regular dict
    return {
        "id":           d["id"],
        "message":      d["message"],
        "final_amt":    d.get("final_amt"),
        "reward_exec":  d.get("reward_exec"),
        "reward_cust":  d.get("reward_cust"),
        "fee_platform": d.get("fee_platform"),
        "status":       d["status"],
        "created_at":   d["created_at"],
        "closed_at":    d.get("updated_at"),
        "driver_fio":   d.get("driver_fio"),
        "loads":        json.loads(d["loads"])   if d.get("loads")   else [],
        "unloads":      json.loads(d["unloads"]) if d.get("unloads") else [],
    }

# --- alias for backward‑compatibility -----------------------------
@app.get("/orders/{order_id}")
async def get_order_alias(order_id: int):
    """
    Shortcut endpoint that simply proxies to /order/{order_id}.
    Needed for old bot callbacks that still reference /orders/…
    """
    return await get_order(order_id)
# ------------------------------------------------------------------


# ---------- ADMIN: общая статистика ----------
@app.get("/admin/overview")
async def admin_overview():
    """Короткая сводка: количество агентов, заявок и прибыль платформы за текущие сутки."""
    # --- количество агентов ---
    cust_agents = await database.fetch_val(
        select(func.count()).select_from(agents).where(agents.c.agent_type == "заказчик")
    ) or 0
    exec_agents = await database.fetch_val(
        select(func.count()).select_from(agents).where(agents.c.agent_type == "исполнитель")
    ) or 0

    # --- всего заявок ---
    orders_total = await database.fetch_val(
        select(func.count()).select_from(orders)
    ) or 0

    # --- прибыль платформы за последние 24 часа ---
    now = datetime.utcnow()
    day_start = now - timedelta(days=1)
    profit_day = await database.fetch_val(
        select(func.coalesce(func.sum(orders.c.fee_platform), 0))
        .where(orders.c.status == "paid")
        .where(orders.c.updated_at >= day_start)
    ) or 0

    return {
        "cust_agents": int(cust_agents),
        "exec_agents": int(exec_agents),
        "orders_total": int(orders_total),
        "profit_day": int(profit_day),
    }


# ---------- ADMIN: списки заявок по статусу ----------

@app.get("/admin/orders")
async def admin_orders(
    status: Optional[str] = None,
    page: int = 0,
    limit: int = 15,
):
    """
    Вернуть заявки определённого статуса постранично (по 15 шт).

    Параметры:
    • status = confirmed | in_progress | done | paid
    • page   = 0,1,2 …   (смещение = page*limit)
    """
    # If no status provided, return all current (non-paid) orders
    if status is None:
        rows = await database.fetch_all(
            orders.select()
                  .where(orders.c.status.in_(CURRENT_STATUSES))
                  .order_by(orders.c.id.desc())
                  .offset(page * limit)
                  .limit(limit)
        )
        return [
            {"id": r["id"], "message": r["message"], "status": r["status"]}
            for r in rows
        ]

    if status not in STATUSES_ALL:
        raise HTTPException(status_code=400, detail="bad status")

    rows = await database.fetch_all(
        orders.select()
              .where(orders.c.status == status)
              .order_by(orders.c.id.desc())
              .offset(page * limit)
              .limit(limit)
    )
    return [
        {
            "id": r["id"],
            "message": r["message"][:60] + ("…" if len(r["message"]) > 60 else "")
        }
        for r in rows
    ]


# ---------- ADMIN: all CURRENT (не оплаченные) заявки ----------
@app.get("/admin/current_orders")
async def admin_current_orders(
    page: int = 0,
    limit: int = 15,
):
    """
    Вернуть все незакрытые (не оплаченные) заявки постранично.

    Эквивалентно status IN ("active", "confirmed", "in_progress").
    Используется кнопкой «Текущие» в админ‑панели бота.
    """
    rows = await database.fetch_all(
        orders.select()
              .where(orders.c.status.in_(CURRENT_STATUSES))
              .order_by(orders.c.id.desc())
              .offset(page * limit)
              .limit(limit)
    )
    return [
        {
            "id": r["id"],
            "message": r["message"][:60] + ("…" if len(r["message"]) > 60 else ""),
            "status": r["status"],
        }
        for r in rows
    ]


# ---------- ADMIN: детальная карточка заявки ----------

@app.get("/admin/order/{order_id}")
async def admin_order(order_id: int):
    """
    Полная карточка заявки для администратора:
    включает обе роли, суммы, driver_fio и т.д.
    """
    row = await database.fetch_one(
        orders.select().where(orders.c.id == order_id)
    )
    if not row:
        raise HTTPException(status_code=404, detail="Order not found")
    row = dict(row)

    # --- закачик ---
    cust = await database.fetch_one(
        agents.select().where(agents.c.telegram_id == row["telegram_id"])
    )
    cust = dict(cust) if cust else {}

    # --- исполнитель ---
    exec_ = None
    if row.get("executor_id"):
        exec_row = await database.fetch_one(
            agents.select().where(agents.c.telegram_id == row["executor_id"])
        )
        exec_ = dict(exec_row) if exec_row else {}

    return {
        "id": row["id"],
        "message": row["message"],
        "status": row["status"],
        "original_amt": row["original_amt"],
        "final_amt": row["final_amt"],
        "reward_cust": row["reward_cust"],
        "reward_exec": row["reward_exec"],
        "fee_platform": row["fee_platform"],
        "driver_fio": row.get("driver_fio"),
        "customer": cust,
        "executor": exec_,
        "created_at": row["created_at"],
        "updated_at": row.get("updated_at"),
    }

# ---------- profile & personal lists ----------

@app.get("/agent/{telegram_id}")
async def get_agent(telegram_id: int):
    """Вернуть профиль агента по Telegram‑ID."""
    row = await database.fetch_one(
        agents.select().where(agents.c.telegram_id == telegram_id)
    )
    if not row:
        raise HTTPException(status_code=404, detail="Agent not found")
    row = dict(row)  # convert Record to plain dict
    return {
        "name": row["name"],
        "agent_type": row["agent_type"],
        "phone": row.get("phone", "")
    }


@app.get("/orders/by_customer/{telegram_id}")
async def orders_by_customer(telegram_id: int, limit: int = 30):
    """
    Все заявки заказчика (без фильтра по статусу). Используется для меню «Мои заявки» у Навигатора.
    """
    rows = await database.fetch_all(
        orders.select()
              .where(orders.c.telegram_id == telegram_id)
              .order_by(orders.c.id.desc())
              .limit(limit)
    )
    return [
        {
            "id":          d["id"],
            "message":     d["message"],
            "status":      d["status"],
            "original_amt": d["original_amt"],
            "final_amt":    d["final_amt"],
            "fee_platform": d["fee_platform"],
            "created_at":   d.get("created_at"),
            "updated_at":   d.get("updated_at"),
        }
        for d in map(dict, rows)   # ← ключевая разница
    ]

@app.get("/orders/by_customer_open/{telegram_id}")
async def orders_by_customer_open(telegram_id: int, limit: int = 30):
    """
    Незакрытые заявки заказчика: только active, confirmed, done (без paid).
    Используется для меню «Все актуальные заявки» у заказчиков.
    """
    rows = await database.fetch_all(
        orders.select()
              .where(orders.c.telegram_id == telegram_id)
              .where(orders.c.status.in_(("active", "confirmed", "in_progress", "done")))
              .order_by(orders.c.id.desc())
              .limit(limit)
    )
    return [
        {
            "id":          d["id"],
            "message":     d["message"],
            "status":      d["status"],
            "original_amt": d["original_amt"],
            "final_amt":    d["final_amt"],
            "reward_cust":  d["reward_cust"],
            "fee_platform": d["fee_platform"],
        }
        for d in map(dict, rows)
    ]

# --- Новый эндпоинт: только закрытые заявки конкретного заказчика ---
@app.get("/orders/by_customer_closed/{telegram_id}")
async def orders_by_customer_closed(telegram_id: int, limit: int = 30):
    """
    Список ЗАКРЫТЫХ заявок, опубликованных данным заказчиком.
    Используется в меню «Мой кабинет» (заказчики).
    """
    rows = await database.fetch_all(
        orders.select()
              .where(orders.c.telegram_id == telegram_id)
              .where(orders.c.status == "paid")
              .order_by(orders.c.updated_at.desc())
              .limit(limit)
    )
    return [
        {
            "id":          d["id"],
            "message":     d["message"],
            "status":      d["status"],
            "original_amt": d["original_amt"],
            "final_amt":    d["final_amt"],
            "reward_cust":  d["reward_cust"],
            "fee_platform": d["fee_platform"],
            "closed_at":    d["updated_at"],
        }
        for d in map(dict, rows)
    ]

@app.get("/orders/by_executor/{telegram_id}")
async def orders_by_executor(telegram_id: int, limit: int = 30):
    """Список заявок, закрытых данным исполнителем (только оплаченные)."""
    rows = await database.fetch_all(
        orders.select()
              .where(orders.c.executor_id == telegram_id)
              .where(orders.c.status == "paid")
              .order_by(orders.c.id.desc())
              .limit(limit)
    )
    return [
        {
            "id":          d["id"],
            "message":     d["message"],
            "status":      d["status"],
            "final_amt":   d["final_amt"],
            "reward_exec": d["reward_exec"],
            "closed_at":   d.get("updated_at"),
            "created_at":  d.get("created_at"),
            "updated_at":  d.get("updated_at"),
            "executor_id": d.get("executor_id"),
        }
        for d in map(dict, rows)      # ← ключевая разница
    ]

# --- Новый эндпоинт: незакрытые (не оплаченные) заявки исполнителя ---
@app.get("/orders/by_executor_open/{telegram_id}")
async def orders_by_executor_open(telegram_id: int, limit: int = 30):
    rows = await database.fetch_all(
        orders.select()
              .where(orders.c.executor_id == telegram_id)
              .where(orders.c.status.in_(("confirmed", "done")))
              .order_by(orders.c.id.desc())
              .limit(limit)
    )
    return [
        {
            "id":          d["id"],
            "message":     d["message"],
            "status":      d["status"],
            "final_amt":   d["final_amt"],
            "reward_exec": d["reward_exec"],
            "updated_at":  d.get("updated_at"),
            "loads":       json.loads(d["loads"])   if d.get("loads")   else [],
            "unloads":     json.loads(d["unloads"]) if d.get("unloads") else [],
        }
        for d in map(dict, rows)
    ]

# ---------- Исполнитель: показать все задачи (role="исполнитель") ----------
@app.get("/tasks/by_executor/{telegram_id}")
async def show_tasks(telegram_id: int, limit: int = 30):
    """
    Список всех задач для исполнителя (роль "исполнитель").
    """
    url = f"/orders/by_executor_open/{telegram_id}"
    rows = await _fetch_orders(url)
    debug_str = "DEBUG (для Драйвера):\n"
    debug_str += f"rows: {rows}\n"
    print(debug_str)
    # await update.message.reply_text(debug_str, parse_mode=ParseMode.MARKDOWN)
    # фильтрация и остальной код без изменений (оставить как есть)
    return rows

# ---------- COMPANY endpoints ----------
@app.post("/company", status_code=201)
async def save_company(item: CompanyIn):
    """
    Создать или обновить компанию по ИНН.

    • ИНН очищается от любых символов, кроме цифр.
    • При конфликте по ИНН существующая запись полностью
      заменяется новым payload (без JSON-merge).
    """
    inn_clean = re.sub(r"\D", "", item.inn or "")
    if not inn_clean:
        raise HTTPException(status_code=400, detail="inn required")

    # формируем payload без unset-полей
    payload = item.dict(exclude_unset=True)

    query = """
    INSERT INTO companies (inn, data)
    VALUES (:inn, CAST(:data AS jsonb))
    ON CONFLICT (inn)
    DO UPDATE
       SET data = EXCLUDED.data
    """

    await database.execute(
        query=query,
        values={
            "inn": inn_clean,
            "data": json.dumps(payload, ensure_ascii=False),
        },
    )
    return {"status": "ok", "inn": inn_clean}


@app.get("/company/{inn}")
async def get_company(inn: str):
    """Return company data by cleaned INN."""
    inn_clean = re.sub(r"\D", "", inn or "")
    row = await database.fetch_one(
        company.select().where(company.c.inn == inn_clean)
    )
    if not row:
        raise HTTPException(status_code=404, detail="Not Found")
    return json.loads(row["data"])

# ---------- ADMIN: create invite token ----------
@app.post("/admin/invite")
async def admin_create_invite(inv: InviteCreate):
    if inv.role not in ("заказчик", "исполнитель"):
        raise HTTPException(status_code=400, detail="bad role")
    token = secrets.token_urlsafe(8)
    expires = datetime.utcnow() + timedelta(hours=inv.ttl_hours)
    invites_db[token] = {"role": inv.role, "expires": expires.isoformat(), "used_by": None}
    save_invites()
    link = f"https://t.me/{BOT_USERNAME}?start={token}"
    return {
        "token": token,
        "expires": expires,
        "role": inv.role,
        "deep_link": link,
    }


# ---------- INVITE: claim token (used by bot) ----------
@app.post("/invite/claim")
async def invite_claim(claim: InviteClaim):
    inv = invites_db.get(claim.token)
    if not inv:
        raise HTTPException(status_code=404, detail="invalid token")
    if inv["used_by"]:
        raise HTTPException(status_code=403, detail="token already used")
    if datetime.fromisoformat(inv["expires"]) < datetime.utcnow():
        raise HTTPException(status_code=403, detail="token expired")

    # mark used
    inv["used_by"] = claim.telegram_id
    inv["used_at"] = datetime.utcnow().isoformat()
    save_invites()
    return {"role": inv["role"]}


# --------- NEW: Leaderboard & Achievements ---------

@app.get("/leaderboard")
async def leaderboard(
    metric: str = "amount",   # amount | deals
    period: str = "week",     # week | month
    limit: int = 10,
):
    """
    Заглушка: топ‑агентов.
    Возвращает список словарей [{rank,name,deals,amount}, …]
    """
    demo = [
        {"rank": 1, "name": "Иван И.",  "deals": 17, "amount": 950_000},
        {"rank": 2, "name": "Мария П.", "deals": 15, "amount": 900_000},
        {"rank": 3, "name": "Тестовый Польз.", "deals": 3, "amount": 120_000},
    ]
    return demo[:limit]


@app.get("/achievements/{telegram_id}")
async def achievements(telegram_id: int):
    """
    Заглушка: достижения и события одного пользователя.
    Возвращает {"events":[...], "badges":[...]}.
    """
    events = [
        {"date": "10.05.25", "type": "level", "text": "Достиг уровня Приют‑11"},
        {"date": "05.05.25", "type": "deal",  "text": "Выполнена заявка #12345"},
        {"date": "01.05.25", "type": "bonus", "text": "Бонус 5 000 ₽ за апрель"},
    ]
    badges = [
        {"type": "elbrus",   "obtained": True},
        {"type": "top_week", "obtained": False},
        {"type": "million",  "obtained": True},
    ]
    return {"events": events, "badges": badges}



# ---------- ADMIN: reset (delete) agent by Telegram‑ID ----------
@app.delete("/admin/reset_agent/{telegram_id}")
async def admin_reset_agent(telegram_id: int):
    """
    Полностью удаляет агента из БД и локального кэша,
    чтобы он мог зарегистрироваться заново по другому токену.
    """
    # 1) удалить из Postgres
    await database.execute(
        agents.delete().where(agents.c.telegram_id == telegram_id)
    )

    # 2) удалить из in‑memory кэша
    agents_db.pop(telegram_id, None)

    # 3) обновить файл agents.json
    save_agents()

    return {"status": "ok", "removed_id": telegram_id}



# ---------- ADMIN: broadcast message to agents ----------
class BroadcastMsg(BaseModel):
    text: str
    role: str | None = None        # 'заказчик' | 'исполнитель' | None
    image_url: str | None = None   # URL или file_id для sendPhoto

@app.post("/admin/broadcast")
async def admin_broadcast(msg: BroadcastMsg):
    """
    Рассылка уведомления всем агентам.
    • text – текст уведомления (обязательно)
    • role – опционально фильтровать: 'заказчик', 'исполнитель', либо null -> всем
    """
    # выбираем адресатов
    if msg.role in ("заказчик", "исполнитель"):
        rows = await database.fetch_all(
            agents.select().where(agents.c.agent_type == msg.role)
        )
    else:
        rows = await database.fetch_all(agents.select())

    sent, errors = 0, 0
    for r in rows:
        try:
            chat_id = int(r["telegram_id"])
            if msg.image_url:           # отправляем картинку + подпись
                url = f"https://api.telegram.org/bot{BOT_TOKEN}/sendPhoto"
                payload = {
                    "chat_id": chat_id,
                    "photo": msg.image_url,
                    "caption": msg.text,
                }
            else:                       # обычный текст
                url = f"https://api.telegram.org/bot{BOT_TOKEN}/sendMessage"
                payload = {
                    "chat_id": chat_id,
                    "text": msg.text,
                }
            requests.post(url, json=payload, timeout=4)
            sent += 1
        except Exception as e:
            errors += 1
            print("Broadcast send error:", e)

    return {"status": "ok", "sent": sent, "errors": errors}

# ---------- ADMIN: create a quick demo order (for terminal testing) ----------
from random import randint, choice
from uuid import uuid4

@app.post("/admin/mock_order")
async def admin_mock_order(
    telegram_id: int = 111,          # любой TG‑ID заказчика
    from_city: str = "Москва",
    to_city: str   = "Казань",
    price: int = 45000,
):
    """
    Создаёт тестовую заявку с двумя точками погрузки и выгрузки
    без валидаций. Удобно прогонять через терминал:

        curl -X POST "http://127.0.0.1:8000/admin/mock_order?telegram_id=555"
    """
    add_payload = {
        "telegram_id": telegram_id,
        "message": f"{from_city} — {to_city} {price} руб",
        "original_amt": price,
        "cust_company_name": f"ООО Тест-{randint(100,999)}",
        "cust_director":     "Петров П.П.",
        "cust_requisites":   f"ИНН 7700{randint(100000,999999)}",
        "vin_list": [f"VIN{randint(1000,9999)}" for _ in range(3)],
        "loads": [
            {"place": f"{from_city}, ул. Ленина 1",
             "date": datetime.utcnow().strftime('%d.%m %H:%M'),
             "contact": "Иван +7-900-000-00-00",
             "vins": []},
        ],
        "unloads": [
            {"place": f"{to_city}, ул. Тукая 5",
             "date": (datetime.utcnow()+timedelta(days=1)).strftime('%d.%m %H:%M'),
             "contact": "Пётр +7-901-000-00-00",
             "vins": []},
        ],
        "pay_terms": "100% безнал с НДС, в течение 3 дней",
    }
    # Re‑use existing function
    resp = await add_order(Order(**add_payload))
    return resp

# ---------- FILE DOWNLOAD helper ----------
@app.get("/file")
async def download_file(path: str):
    """
    Скачивание файла по абсолютному или относительному пути.

    Бот вызывает:  GET /file?path=acts/act_16_exec.docx
    Если путь относительный — дополнительно проверяем внутри BASE_DIR.
    """
    # если путь относительный — интерпретируем относительно BASE_DIR
    if not os.path.isabs(path):
        path = os.path.join(BASE_DIR, path)

    if not os.path.isfile(path):
        raise HTTPException(status_code=404, detail="File not found")

    return FileResponse(path, filename=os.path.basename(path))